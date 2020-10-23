# -*- coding: utf-8 -*-
"""
Created on Tue Jul 14 14:35:26 2020

@author: panglei
"""
import os
import sys
import seaborn as sns
import pandas as pd
from pandas import Series, DataFrame
import numpy as np
import matplotlib.pyplot as plt
import matplotlib as mpl
from datetime import *  
from datetime import datetime
from pandas.tseries.offsets import Day,MonthEnd
import matplotlib.dates as mdate
import time

#获取时间
curr_time = datetime.now()
time_str = (curr_time.date()-Day()).strftime("%Y-%m-%d")
time_str_1 = (curr_time.date()-Day()).strftime("%Y%m%d")
time_str_30 = (curr_time.date()-32*Day()).strftime("%Y%m%d")
#改变路径
#TODO
df_old_mchnt_dir = r'商户交易累计/2020-06-03-10-18 商户交易.xlsx'
df_new_mchnt_dir = r'商户交易每日/2020-10-20 商户交易.txt'
df_mchnt_dir = r"C:/工作/典型事件/手机POS交易数据疑似套现/拉卡拉商户交易明细/商户交易累计/2020-06-03-10-20 商户交易.xlsx"
document_dir = r'C:/工作/典型事件/手机POS交易数据疑似套现/拉卡拉商户交易明细/商户交易日监测报告/商户交易日监测报告'+time_str+'.docx'
pic_dir = r'C:/工作/典型事件/手机POS交易数据疑似套现/拉卡拉商户交易明细/图片/'
#设置画图字体
mpl.rcParams['font.sans-serif'] = ['SimHei']
mpl.rcParams['font.serif'] = ['MSYH.TTC']
mpl.rcParams['axes.unicode_minus'] = False # 解决保存图像是负号'-'显示为方块的问题,或者转换负号为字符串
output_text = []
#柱形图加标签
def show_value_for_barplot(barplot, h_v="v",form='.2f',percent=1):
    if h_v == "v":
        for p in barplot.patches:
            per = str(format(p.get_height()*percent,form))+'%'
            barplot.annotate(per,
                             (p.get_x() + p.get_width() / 2., p.get_height()),
                             ha = 'center', va = 'center', xytext = (0, 5),
                             textcoords = 'offset points')
    elif h_v == "h":
        for p in barplot.patches:
            # @param format(p.get_width(), '.2f'), word in string format you want to put in the figure
            # @param (p.get_width(), p.get_y()+ p.get_height() / 2.), x and y pos of word
            # @param xytext, offset of word
            barplot.annotate(format(p.get_width(), '.2f'),
                             (p.get_width(), p.get_y()+ p.get_height() / 2.),
                             ha = 'center', va = 'center', xytext = (30, 0),
                             textcoords = 'offset points')
#导入原始数据
os.chdir(r'C:/工作/典型事件/手机POS交易数据疑似套现/拉卡拉商户交易明细')

df_old = pd.read_excel(df_old_mchnt_dir,sheet_name = 'Sheet1',dtype=object,header=0)
df_new_colname=['pri_acct_no_conv_sm3', 'card_attr', 'iss_ins_id_cd', 'acpt_ins_id_cd',
                'fwd_ins_id_cd', 'loc_trans_tm', 'hp_settle_dt', 'mchnt_cd',
                'card_accptr_nm_addr', 'trans_at', 'mchnt_tp', 'term_id', 'trans_chnl',
                'trans_id', 'pos_entry_md_cd', 'sys_tra_no', 'resp_cd4',
                'ext_hce_prod_nm', 'ext_hce_prod_in', 'ext_conn_in', 'record_dt','sys_record_dt']
df_new = pd.read_table(df_new_mchnt_dir,delimiter=',',
                 header=0, squeeze=True,dtype=object,names=df_new_colname)
del df_new['sys_record_dt']
df_old['trans_at'] = pd.to_numeric(df_old['trans_at'], errors='coerce').fillna(0)
df_new['trans_at'] = (pd.to_numeric(df_new['trans_at'], errors='coerce').fillna(0))/100
# Index(['pri_acct_no_conv_sm3', 'card_attr', 'iss_ins_id_cd', 'acpt_ins_id_cd',
#        'fwd_ins_id_cd', 'loc_trans_tm', 'hp_settle_dt', 'mchnt_cd',
#        'card_accptr_nm_addr', 'trans_at', 'mchnt_tp', 'term_id', 'trans_chnl',
#        'trans_id', 'pos_entry_md_cd', 'sys_tra_no', 'resp_cd4',
#        'ext_hce_prod_nm', 'ext_hce_prod_in', 'ext_conn_in', 'record_dt'],
#       dtype='object')
df = pd.concat([df_old,df_new], axis=0)
#去重
#subset=None,keep='first',inplace=False
df = df.drop_duplicates(subset=df_new_colname[:-2],keep='first')
writer = pd.ExcelWriter(df_mchnt_dir)
df.to_excel(writer, index=False,encoding='utf-8',sheet_name='Sheet1')
writer.save()

sucess_list = ['00','']
df_sucess = df[df['resp_cd4'].isin(sucess_list)]
#df_sucess['trans_at'] = pd.to_numeric(df_sucess['trans_at'], errors='coerce').fillna(0)

#总交易金额
trans_at_sum = df_sucess['trans_at'].sum()
#总交易笔数
trans_num_sum = len(df_sucess['pri_acct_no_conv_sm3'])
#总商户数量
mchnt_sum = len(df_sucess['mchnt_cd'].unique())
#总卡片数量
card_sum = len(df_sucess['pri_acct_no_conv_sm3'].unique())

#%%交易金额随时间变化趋势
df_sucess_time = df_sucess
df_sucess_time['hp_settle_dt'] = pd.to_datetime(df_sucess['hp_settle_dt'])
df_sucess_time.index = df_sucess_time['hp_settle_dt']
df_sucess_time.rename(columns={'hp_settle_dt':'hp_settle_dt_origin'}, inplace = True)
#del df_sucess_time['hp_settle_dt']
df_sucess_time = df_sucess_time.sort_index()
trans_at_day = df_sucess_time.groupby('hp_settle_dt')['trans_at'].agg('sum')
plt.figure(dpi=600)#设置分辨率
#splot,ax = trans_at_day.plot()
i=0
for a,b in zip(trans_at_day.index,trans_at_day.values):
    i=i+1
    if i%8==0:
        plt.text(a, b+80000, '%.2f' % (b/10000), ha='center', va= 'bottom',fontsize=6,
                 bbox=dict(boxstyle="round",ec=(0.9,1., 0.9),fc=(0.9,1., 0.9)))
    elif i%4==0:
        plt.text(a, b-80000, '%.2f' % (b/10000), ha='center', va= 'bottom',fontsize=6,
                 bbox=dict(boxstyle="round",ec=(0.9,1., 0.9),fc=(0.9,1., 0.9)))

ax = trans_at_day.plot()
ax.xaxis.set_minor_locator(mdate.WeekdayLocator(byweekday=(1),
                                                interval=1))
ax.xaxis.set_minor_formatter(mdate.DateFormatter('%d'))
ax.xaxis.set_major_locator(mdate.MonthLocator())
ax.xaxis.set_major_formatter(mdate.DateFormatter('\n%b\n%Y'))
plt.tight_layout()
plt.xticks(rotation=0)#设置刻度旋转角度
plt.xlabel('时间',fontsize=11)#设置刻度标签
plt.ylabel('交易金额(元)',fontsize=11)
plt.grid(axis="y",linewidth=0.5)#color='r',linestyle='-.'
plt.savefig(pic_dir+"交易金额随时间变化曲线.png",bbox_inches = 'tight')

#%%日活跃商户随时间变化趋势
def func_mchnt_day(list_mchnt):
    return list_mchnt.nunique()
mchnt_at_day = df_sucess_time.groupby('hp_settle_dt')['mchnt_cd'].apply(func_mchnt_day)
plt.figure(dpi=600)#设置分辨率
splot = mchnt_at_day.plot()
i=0
for a,b in zip(mchnt_at_day.index,mchnt_at_day.values):
    i=i+1
    if i%6==0:
        plt.text(a, b+30, '%.0f' % (b), ha='center', va= 'bottom',fontsize=6,
                 bbox=dict(boxstyle="round",ec=(0.9,1., 0.9),fc=(0.9,1., 0.9)))
    elif i%3==0:
        plt.text(a, b-40, '%.0f' % (b), ha='center', va= 'bottom',fontsize=6,
                 bbox=dict(boxstyle="round",ec=(0.9,1., 0.9),fc=(0.9,1., 0.9)))
plt.xticks(rotation=15)#设置刻度旋转角度
plt.xlabel('时间',fontsize=11)#设置刻度标签
plt.ylabel('活跃商户数量',fontsize=11)
plt.grid(axis="y",linewidth=0.5)#color='r',linestyle='-.'
plt.savefig(pic_dir+"日活跃商户随时间变化趋势.png",bbox_inches = 'tight')
#%%交易笔数随时间变化趋势
plt.figure(dpi=600)#设置分辨率
#splot = df_sucess_time.groupby('hp_settle_dt')['trans_at'].agg('count').plot()
m=df_sucess_time.groupby('hp_settle_dt')['trans_at'].agg('count')
m.plot()
i=0
for a,b in zip(m.index,m.values):
    i=i+1
    if i%6==0:
        plt.text(a, b+20, '%.0f' % b, ha='center', va= 'bottom',fontsize=7,
                 bbox=dict(boxstyle="round",ec=(1., 0.9, 0.9),fc=(1., 0.9, 0.9)))
plt.xticks(rotation=15)#设置刻度旋转角度
plt.xlabel('时间',fontsize=11)#设置刻度标签
plt.ylabel('交易笔数',fontsize=11)
plt.grid(axis="y",linewidth=0.5)#color='r',linestyle='-.'
plt.savefig(pic_dir+"交易笔数随时间变化曲线.png",bbox_inches = 'tight')
#笔均金额
plt.figure(dpi=600)#设置分辨率
trans_at_avg_bill = df_sucess_time.groupby('hp_settle_dt')['trans_at'].agg('mean')
trans_at_avg_bill.plot()
i=0
for a,b in zip(trans_at_avg_bill.index,trans_at_avg_bill.values):
    i=i+1
    if i%6==0:
        plt.text(a, b+200, '%.0f' % b, ha='center', va= 'bottom',fontsize=7,
                 bbox=dict(boxstyle="round",ec=(1., 0.9, 0.9),fc=(1., 0.9, 0.9)))
        #plt.annotate('%.0f' % b, xy=(a, b), arrowprops=dict(arrowstyle='->'),fontsize=7)
        #+random.randint(0,399)
plt.xticks(rotation=15)#设置刻度旋转角度
plt.xlabel('时间',fontsize=11)#设置刻度标签
plt.ylabel('笔均金额（元）',fontsize=11)
plt.grid(axis="y",linewidth=0.5)#color='r',linestyle='-.'
plt.savefig(pic_dir+"笔均金额随时间变化曲线.png",bbox_inches = 'tight')
#ha有三个选择：right,center,left
#va有四个选择：'top', 'bottom', 'center', 'baseline'
#%%刷卡与二维码比例
#方法一
# QR_list = ['042','041','942']
# df_sucess_QR = df_sucess[df_sucess['pos_entry_md_cd'].isin(QR_list)]
# df_sucess_card = df_sucess[~df_sucess['pos_entry_md_cd'].isin(QR_list)]
#方法二
QR_list = ['Y','Z']
QR_list_1 = ['Y','Z','0','5']
df_sucess_QR = df_sucess[df_sucess['ext_hce_prod_nm'].isin(QR_list)]
df_sucess_quickpass = df_sucess[~df_sucess['ext_hce_prod_nm'].isin(QR_list_1)]
df_sucess_card = df_sucess[df_sucess['ext_hce_prod_nm'].isin(['0'])]

pos_entry_md_cd_distr = Series([len(df_sucess_QR),len(df_sucess_quickpass),len(df_sucess_card)],
                     index=['二维码','手机闪付','刷卡'])
plt.figure(dpi=600)#设置分辨率
sns.set_palette("muted")
pos_entry_md_cd_distr.plot.pie(autopct='%.2f%%',pctdistance = 0.85,startangle=90,
                                explode = [0, 0.1, 0],wedgeprops = {'width': 0.4,'edgecolor': 'w'})
plt.xlabel('刷卡与二维码笔数比例',fontsize=11)#设置刻度标签
plt.ylabel('',fontsize=11)
plt.savefig(pic_dir+"刷卡与二维码笔数比例.png",bbox_inches = 'tight')

pos_entry_md_cd_transat_distr = Series([df_sucess_QR['trans_at'].sum(),
                                        df_sucess_quickpass['trans_at'].sum(),df_sucess_card['trans_at'].sum()],
                                        index=['二维码','手机闪付','刷卡'])
plt.figure(dpi=600)#设置分辨率
sns.set_palette("RdBu")
pos_entry_md_cd_transat_distr.plot.pie(autopct='%.2f%%',pctdistance = 0.85,startangle=90,
                               explode = [0, 0.1, 0],wedgeprops = {'width': 0.4,'edgecolor': 'w'})
plt.xlabel('刷卡与二维码交易金额比例',fontsize=11)#设置刻度标签
plt.ylabel('',fontsize=11)
plt.savefig(pic_dir+"刷卡与二维码金额比例.png",bbox_inches = 'tight')


#%%金额段分布
customer = ['S22','S56','S46','S10','S65','S48','S20','S35','S67','S49','S50','W20','W21','']
df_trans_at = df_sucess[df_sucess['trans_id'].isin(customer)]
sections = Series([-1,0,100,1000,2500,5000,35000,99999999999999])
group_names = ['0','0~100','100~1000','1000~2500','2500~5000','5000~35000','35000以上']
cuts = pd.cut(df_trans_at['trans_at'],sections,labels=group_names)
df_trans_at_order = cuts.value_counts().sort_values(ascending=False)
plt.figure(dpi=600)#设置分辨率
splot = cuts.value_counts().plot(kind='bar')
show_value_for_barplot(splot,h_v="v",form = '.2f',percent=100/trans_num_sum)
plt.xticks(rotation=0)#设置刻度旋转角度
plt.xlabel('金额',fontsize=11)#设置刻度标签
plt.ylabel('交易笔数',fontsize=11)
plt.title('交易金额消费区间段分布')
plt.savefig(pic_dir+"消费金额区间段分布.png",bbox_inches = 'tight')
plt.show()

output_text.append('金额区间段分布(消费)最多的前三个为：'+
                       str(df_trans_at_order.index[0])+'、'+
                       str(df_trans_at_order.index[1])+'、'+
                       str(df_trans_at_order.index[2])+'、'+
                       ',其中最密集金额区间段交易占总交易笔数的'+str(splot.patches[0].get_height()*100/trans_num_sum)+'%')
#%%2000/5000
#二维码
trans_at_QR_cut = Series([df_sucess_QR[df_sucess_QR['trans_at']<2000]['trans_at'].count(),
                          df_sucess_QR[(df_sucess_QR['trans_at']>=2000)|(df_sucess_QR['trans_at']<=5000)]['trans_at'].count(),
                          df_sucess_QR[df_sucess_QR['trans_at']>5000]['trans_at'].count()],
                         index=['<2000扫码总笔数','>2000&<5000扫码总笔数','>5000扫码总笔数'])
plt.figure(dpi=600)#设置分辨率
sns.set_palette("RdBu")
trans_at_QR_cut.plot.pie(autopct='%.2f%%',pctdistance = 0.85,startangle=90,
                         explode = [0, 0.1, 0],wedgeprops = {'width': 0.4,'edgecolor': 'w'})
plt.xlabel('二维码交易金额区间笔数',fontsize=11)#设置刻度标签
plt.ylabel('',fontsize=11)
plt.savefig(pic_dir+"二维码交易金额区间笔数.png",bbox_inches = 'tight')
#刷卡
trans_at_card_cut = Series([df_sucess_card[df_sucess_card['trans_at']<2000]['trans_at'].count(),
                          df_sucess_card[(df_sucess_card['trans_at']>=2000)|(df_sucess_card['trans_at']<=5000)]['trans_at'].count(),
                          df_sucess_card[df_sucess_card['trans_at']>5000]['trans_at'].count()],
                         index=['<2000刷卡总笔数','>2000&<5000刷卡总笔数','>5000刷卡总笔数'])
plt.figure(dpi=600)#设置分辨率
sns.set_palette("RdBu")
trans_at_card_cut.plot.pie(autopct='%.2f%%',pctdistance = 0.85,startangle=90,
                         explode = [0, 0.1, 0],wedgeprops = {'width': 0.4,'edgecolor': 'w'})
plt.xlabel('刷卡交易金额区间笔数',fontsize=11)#设置刻度标签
plt.ylabel('',fontsize=11)
plt.savefig(pic_dir+"刷卡交易金额区间笔数.png",bbox_inches = 'tight')                      
#%%借贷记分布
df_single_card = df_sucess.drop_duplicates(['pri_acct_no_conv_sm3'])
plt.figure(dpi=600)#设置分辨率
card_attr_distr = df_single_card['card_attr'].value_counts()
#splot = card_attr_distr.plot(kind='bar')
card_attr_distr = card_attr_distr.rename({'01': '借记卡', '02': '贷记卡', '03': '准贷记卡'}, axis='index')
#show_value_for_barplot(splot,h_v="v",percent=100/card_sum)
card_attr = Series(['其他','借记卡','贷记卡','准贷记卡','借贷合一卡','预付费卡','单用途预付费卡'],
                     index=['0','01','02','03','04','05','06'])
splot = card_attr_distr.plot.pie(autopct='%.2f%%',pctdistance = 0.85,startangle=0,
                               explode = [0, 0.1, 0],wedgeprops = {'width': 0.4,'edgecolor': 'w'})
#output_text.append('卡片最多为：'+str(card_attr[card_attr_distr.index[0]])+
#                   ',数量占总卡片数量的'+str(splot.patches[0].get_height()*100/card_sum)+'%')

#card_attr_distr.plot.pie(autopct='%.2f%%')
plt.ylabel('',fontsize=11)#设置刻度标签
plt.xlabel('借贷记分布',fontsize=11)#设置刻度标签
plt.savefig(pic_dir+"借贷记分布.png",bbox_inches = 'tight')

#%%头部商户情况
#交易金额
mchnt_trans_at = df_sucess.groupby('mchnt_cd')['trans_at'].agg('sum')
mchnt_trans_at = mchnt_trans_at.sort_values(ascending=False)

#交易笔数
mchnt_trans_num = df_sucess.groupby('mchnt_cd')['sys_tra_no'].agg('count')
mchnt_trans_num = mchnt_trans_num.sort_values(ascending=False)

#笔均金额
grouped = df_sucess.groupby('mchnt_cd')
def average_100(arr):
    total = 0
    count = 0
    for element in arr:
        if element>100:
            total = total + element
            count = count + 1
    if count==0:
        return 0
    else:
        return total/count            
trans_at_average_100 = grouped['trans_at'].agg(average_100)
trans_at_average_100 = pd.DataFrame({'mchnt_cd':trans_at_average_100.index,'trans_at_average_100':trans_at_average_100.values})

#日均金额
grouped = df_sucess.groupby('mchnt_cd')
def average_day(df):
    df = df.sort_values(by='hp_settle_dt_origin')
    trans_days = (df['hp_settle_dt_origin'][-1]-df['hp_settle_dt_origin'][0]).days
    if trans_days==0:
        return df['trans_at'].sum()
    else:
        return df['trans_at'].sum()/trans_days
    
trans_at_average_day = grouped[['trans_at','hp_settle_dt_origin']].apply(average_day)
trans_at_average_day = pd.DataFrame({'mchnt_cd':trans_at_average_day.index,'trans_at_average_day':trans_at_average_day.values})

#贷记卡金额占比
grouped = df_sucess[df_sucess['card_attr'].isin(['02','03'])].groupby('mchnt_cd')
trans_num_permachnt_loan = grouped['trans_at'].agg('sum')
trans_num_permachnt_loan = pd.DataFrame({'mchnt_cd':trans_num_permachnt_loan.index,'trans_num_permachnt_loan':trans_num_permachnt_loan.values})


#匹配合并
machnt = pd.merge(mchnt_trans_at,mchnt_trans_num,how='left',on = 'mchnt_cd')
machnt = pd.merge(machnt,trans_at_average_100,how='left',on = 'mchnt_cd')
machnt = pd.merge(machnt,trans_at_average_day,how='left',on = 'mchnt_cd')
machnt = pd.merge(machnt,trans_num_permachnt_loan,how='left',on = 'mchnt_cd')
machnt = pd.merge(machnt,pd.concat([df['mchnt_cd'],df['card_accptr_nm_addr']], axis=1),how='left',on = 'mchnt_cd')
machnt['loan_at_ratio'] = machnt['trans_num_permachnt_loan']/machnt['trans_at']
machnt = machnt.drop_duplicates()

#疑似风险
machnt_risk = machnt[machnt['trans_at']>50000]
machnt_risk = machnt_risk[machnt_risk['sys_tra_no']>10]
#贷记卡交易占比
machnt_risk = machnt_risk[machnt_risk['loan_at_ratio']>0.9]
machnt_risk = machnt_risk[machnt_risk['trans_at_average_day']>10000]
machnt_risk = machnt_risk[machnt_risk['trans_at_average_100']>2000]

df_sended = pd.read_excel(r'C:/工作/典型事件/手机POS交易数据疑似套现/拉卡拉商户交易明细/向拉卡拉报送商户代码台账.xlsx',
                          header=0, squeeze=True,dtype=object)
machnt_risk = machnt_risk[~machnt_risk['mchnt_cd'].isin(df_sended['商户代码'])]
#%%交易金额前十的商户的日交易趋势，从第一笔交易开始。
# for ele in machnt_risk['mchnt_cd'][:5]:
#     plt.figure(dpi=600)#设置分辨率
#     splot = (df_sucess_time[df_sucess_time['mchnt_cd']==ele]).groupby('hp_settle_dt')['trans_at'].agg('sum').plot()
#     plt.xticks(rotation=15)#设置刻度旋转角度
#     plt.xlabel('时间',fontsize=11)#设置刻度标签
#     plt.ylabel('交易金额',fontsize=11)
#     plt.savefig(pic_dir+ele+"交易金额随时间变化曲线.png",bbox_inches = 'tight')

#%%突增/异动/异常情况
#每天的变化情况
def ext_hce_QR(df):
    df_QR = df[df['ext_hce_prod_nm'].isin(QR_list)]
    return df_QR['trans_at'].sum()
def ext_hce_quickpass(df):
    df_QR = df[~df['ext_hce_prod_nm'].isin(QR_list_1)]
    return df_QR['trans_at'].sum()
def ext_hce_card(df):
    df_QR = df[df['ext_hce_prod_nm'].isin(['0'])]
    return df_QR['trans_at'].sum()
sns.set_palette("muted")    
plt.figure(dpi=600)#设置分辨率
splot = df_sucess_time.groupby('hp_settle_dt')[['ext_hce_prod_nm','trans_at']].apply(ext_hce_QR).plot()
splot = df_sucess_time.groupby('hp_settle_dt')[['ext_hce_prod_nm','trans_at']].apply(ext_hce_quickpass).plot( )
splot = df_sucess_time.groupby('hp_settle_dt')[['ext_hce_prod_nm','trans_at']].apply(ext_hce_card).plot()
plt.legend(['二维码','手机闪付','刷卡'])
plt.xticks(rotation=15)#设置刻度旋转角度
plt.xlabel('时间',fontsize=11)#设置刻度标签
plt.ylabel('交易金额',fontsize=11)
plt.savefig(pic_dir+"支付方式交易金额随时间变化曲线.png",bbox_inches = 'tight')

#%%笔均交易金额随时间变化
def ext_hce_QR_mean(df):
    df_QR = df[df['ext_hce_prod_nm'].isin(QR_list)]
    return df_QR['trans_at'].mean()
def ext_hce_quickpass_mean(df):
    df_QR = df[~df['ext_hce_prod_nm'].isin(QR_list_1)]
    return df_QR['trans_at'].mean()
def ext_hce_card_mean(df):
    df_QR = df[df['ext_hce_prod_nm'].isin(['0'])]
    return df_QR['trans_at'].mean()
sns.set_palette("muted")    
plt.figure(dpi=600)#设置分辨率
splot = df_sucess_time.groupby('hp_settle_dt')[['ext_hce_prod_nm','trans_at']].apply(ext_hce_QR_mean).plot()
splot = df_sucess_time.groupby('hp_settle_dt')[['ext_hce_prod_nm','trans_at']].apply(ext_hce_quickpass_mean).plot( )
splot = df_sucess_time.groupby('hp_settle_dt')[['ext_hce_prod_nm','trans_at']].apply(ext_hce_card_mean).plot()
plt.legend(['二维码','手机闪付','刷卡'])
plt.xticks(rotation=15)#设置刻度旋转角度
plt.xlabel('时间',fontsize=11)#设置刻度标签
plt.ylabel('笔均交易金额',fontsize=11)
plt.savefig(pic_dir+"支付方式交易笔均金额随时间变化曲线.png",bbox_inches = 'tight')
#%%单卡不同的手机POS商户情况
#单卡涉及商户数
def single_card_mchnt_func(arr):
    return len(arr.unique())
single_card_mchnt = df_sucess.groupby('pri_acct_no_conv_sm3')['mchnt_cd'].agg(single_card_mchnt_func)
df_single_card_mchnt = pd.DataFrame({'加密卡号':single_card_mchnt.index,'涉及商户数':single_card_mchnt.values})
df_single_card_mchnt_imp_card = df_single_card_mchnt[df_single_card_mchnt['涉及商户数']>=3]

#单卡涉及交易笔数
single_card_trans_num = df_sucess.groupby('pri_acct_no_conv_sm3')['sys_tra_no'].agg('count')
#单卡涉及交易金额
single_card_trans_at = df_sucess.groupby('pri_acct_no_conv_sm3')['trans_at'].agg('sum')

#%%堆叠柱状图
# p1 = plt.bar(ind, S, width, color='#d62728')
# p2 = plt.bar(ind, C, width, bottom=S)
# p3 = plt.bar(ind, M, width, bottom=d)

#%%当日分析
#新增商户数、交易笔数、交易金额
df_curr_day = df[df['record_dt']==time_str_1]
mchnt_curr_day = len(df_curr_day['mchnt_cd'].unique())
df_curr_day_2 = df_sucess_time[df_sucess_time['hp_settle_dt_origin']==time_str]
trans_num_curr_day = df_curr_day_2['sys_tra_no'].count()
trans_at_curr_day = df_curr_day_2['trans_at'].sum()
trans_at_loan_curr_day = ((df_curr_day_2[df_curr_day_2['card_attr'].isin(['02','03'])])['trans_at'].sum())/trans_at_curr_day

#当日可疑商户、历史可疑商户
#交易金额
mchnt_trans_at_curr_day = df_curr_day.groupby('mchnt_cd')['trans_at'].agg('sum')
mchnt_trans_at_curr_day = mchnt_trans_at_curr_day.sort_values(ascending=False)
#交易笔数
mchnt_trans_num_curr_day = df_curr_day.groupby('mchnt_cd')['sys_tra_no'].agg('count')
mchnt_trans_num_curr_day = mchnt_trans_num_curr_day.sort_values(ascending=False)
machnt_curr_day = pd.merge(mchnt_trans_at_curr_day,mchnt_trans_num_curr_day,how='left',on = 'mchnt_cd')
#当日危险商户
machnt_risk_curr_day = machnt_curr_day[machnt_curr_day['trans_at']>15000]
machnt_risk_curr_day = machnt_risk_curr_day[machnt_risk_curr_day['sys_tra_no']>10]
#%%统计活跃商户
#日活
df_day_live = df[df['hp_settle_dt']==time_str_1]
mchnt_day_live = df_day_live['mchnt_cd'].nunique()
#周活
time_str_week_left = (curr_time.date()-8*Day()).strftime("%Y%m%d")
time_str_week_right = (curr_time.date()).strftime("%Y%m%d")
df_week_live = df[(df['hp_settle_dt']>time_str_week_left)&(df['hp_settle_dt']<time_str_week_right)]
mchnt_week_live = df_week_live['mchnt_cd'].nunique()
#月活
time_str_month_left = (curr_time.date()-31*Day()).strftime("%Y%m%d")
time_str_month_right = (curr_time.date()).strftime("%Y%m%d")
df_month_live = df[(df['hp_settle_dt']>time_str_month_left)&(df['hp_settle_dt']<time_str_month_right)]
mchnt_month_live = df_month_live['mchnt_cd'].nunique()

#%%计算过去30天的均值
#交易金额
df_past_30_day = (df_sucess_time[(df_sucess_time['hp_settle_dt_origin']<time_str)&(df_sucess_time['hp_settle_dt_origin']>time_str_30)])
trans_at_30_day = df_past_30_day['trans_at'].sum()
trans_at_30_ratio = trans_at_curr_day/trans_at_30_day*30
#交易笔数
trans_num_30_day = df_past_30_day['sys_tra_no'].count()
trans_num_30_ratio = trans_num_curr_day/trans_num_30_day*30
#贷记卡交易总金额
df_curr_day_loan = (df_curr_day_2[df_curr_day_2['card_attr'].isin(['02','03'])])
df_loan_past_30_day = df_past_30_day[df_past_30_day['card_attr'].isin(['02','03'])]
loan_trans_at_30_day = df_loan_past_30_day['trans_at'].sum()
loan_trans_at_30_ratio = (df_curr_day_loan['trans_at'].sum())/loan_trans_at_30_day*30
#贷记卡交易总笔数
loan_trans_num_30_day = df_loan_past_30_day['sys_tra_no'].count()
loan_trans_num_30_ratio = (df_curr_day_loan['sys_tra_no'].count())/loan_trans_num_30_day*30

#%%5000元以上的贷记卡交易笔数
loan_trans_5000_num_30_day = (df_loan_past_30_day[df_loan_past_30_day['trans_at']>=5000])['sys_tra_no'].count()
loan_trans_5000_num_30_day_ratio = ((df_curr_day_loan[df_curr_day_loan['trans_at']>=5000])['sys_tra_no'].count())/loan_trans_5000_num_30_day*30

#5000元以上的贷记卡交易金额
loan_trans_5000_at_30_day = (df_loan_past_30_day[df_loan_past_30_day['trans_at']>=5000])['trans_at'].sum()
loan_trans_5000_at_30_day_ratio = ((df_curr_day_loan[df_curr_day_loan['trans_at']>=5000])['trans_at'].sum())/loan_trans_5000_at_30_day*30

#%%5000元以上贷记卡交易随时间变化
df_loan_sucess_time_5000 = df_sucess_time[df_sucess_time['trans_at']>=5000]
loan_trans_at_day = df_loan_sucess_time_5000.groupby('hp_settle_dt')['trans_at'].agg('sum')

plt.figure(dpi=600)#设置分辨率
splot = loan_trans_at_day.plot()
i=0
for a,b in zip(loan_trans_at_day.index,loan_trans_at_day.values):
    i=i+1
    if i%6==0:
        plt.text(a, b+20000, '%.2f' % (b/10000), ha='center', va= 'bottom',fontsize=6,
                 bbox=dict(boxstyle="round",ec=(0.9,1., 0.9),fc=(0.9,1., 0.9)))
    elif i%3==0:
        plt.text(a, b-20000, '%.2f' % (b/10000), ha='center', va= 'bottom',fontsize=6,
                 bbox=dict(boxstyle="round",ec=(0.9,1., 0.9),fc=(0.9,1., 0.9)))
plt.xticks(rotation=15)#设置刻度旋转角度
plt.xlabel('时间',fontsize=11)#设置刻度标签
plt.ylabel('交易金额(元)',fontsize=11)
plt.grid(axis="y",linewidth=0.5)#color='r',linestyle='-.'
plt.savefig(pic_dir+"贷记卡5000元以上交易金额随时间变化曲线.png",bbox_inches = 'tight')
#%%
loan_trans_num_day = df_loan_sucess_time_5000.groupby('hp_settle_dt')['trans_at'].agg('count')
plt.figure(dpi=600)#设置分辨率
splot = loan_trans_num_day.plot()
i=0
for a,b in zip(loan_trans_num_day.index,loan_trans_num_day.values):
    i=i+1
    if i%6==0:
        plt.text(a, b+2, '%.0f' % (b), ha='center', va= 'bottom',fontsize=6,
                 bbox=dict(boxstyle="round",ec=(0.9,1., 0.9),fc=(0.9,1., 0.9)))
    elif i%3==0:
        plt.text(a, b-2, '%.0f' % (b), ha='center', va= 'bottom',fontsize=6,
                 bbox=dict(boxstyle="round",ec=(0.9,1., 0.9),fc=(0.9,1., 0.9)))
plt.xticks(rotation=15)#设置刻度旋转角度
plt.xlabel('时间',fontsize=11)#设置刻度标签
plt.ylabel('交易笔数',fontsize=11)
plt.grid(axis="y",linewidth=0.5)#color='r',linestyle='-.'
plt.savefig(pic_dir+"贷记卡5000元以上交易笔数随时间变化曲线.png",bbox_inches = 'tight')
##############################################################################
#%%ump要素上送率
df_uMP_curr_day = df_curr_day_2[df_curr_day_2['term_id'].str[:3]=='uMP']
uMP_upload_ratio = len(df_uMP_curr_day)/len(df_curr_day_2)




#%%输出结果
'''
from  docx import  Document
from  docx.shared import  Pt
from  docx.oxml.ns import  qn
from docx.shared import Inches
os.chdir(r'C:/工作/典型事件/手机POS交易数据疑似套现/拉卡拉商户交易明细')
document = Document()
document.styles['Normal'].font.name = u'仿宋_GB2312'
document.styles['Normal'].font.size=Pt(16)
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
document.add_heading('手机POS监测日报',0)
document.add_heading(u'当日交易情况',1)
#document.add_heading(u'二级标题',2)
picture_width = 4.5
paragraph = document.add_paragraph(time_str+'新增商户'+str(mchnt_curr_day)+'家，新增交易笔数'+
                                   str(trans_num_curr_day)+'笔'+'，新增交易金额'+
                                   str(format(trans_at_curr_day,'.2f'))+'元，'+
                                   '本日活跃商户数为：'+str(mchnt_day_live)+
                                   '本周活跃商户数为：'+str(mchnt_week_live)+
                                   '本月活跃商户数为：'+str(mchnt_month_live)+
                                   '当日疑似危险商户为：'+str(machnt_risk_curr_day))

document.add_heading(u'历史交易情况',1)
paragraph = document.add_paragraph(u'自2020年6月3日以来，手机POS的总交易金额为：'+
                                   str(trans_at_sum)+'元,总交易笔数为：'+
                                   str(trans_num_sum)+'笔,总商户数量为：'+
                                   str(mchnt_sum)+'家,总卡片数量为：'+
                                   str(card_sum)+'张。')
paragraph = document.add_paragraph(u'从历史交易来看，疑似风险商户有以下几家：'+
                                   str(machnt_risk))

paragraph = document.add_paragraph(u'手机POS的交易金额每天的变化趋势如下图：')
document.add_picture('交易金额随时间变化曲线.png',width=Inches(picture_width+0.5))
paragraph = document.add_paragraph(u'手机POS的交易笔数每天的变化趋势如下图：')
document.add_picture('交易笔数随时间变化曲线.png',width=Inches(picture_width+0.5))
paragraph = document.add_paragraph(u'其中涉及的卡片的借贷记分布如下图：')
document.add_picture('借贷记分布.png',width=Inches(picture_width-0.5))
paragraph = document.add_paragraph(u'其中交易金额区间段分布如下图：')
document.add_picture('消费金额区间段分布.png',width=Inches(picture_width+0.5))
paragraph = document.add_paragraph(u'笔均金额随时间变化曲线如下图：')
document.add_picture('笔均金额随时间变化曲线.png',width=Inches(picture_width+0.5)) 
paragraph = document.add_paragraph(u'其中支付方式分布如下图：')
document.add_picture('刷卡与二维码金额比例.png',width=Inches(picture_width-0.5))
document.add_picture('刷卡与二维码笔数比例.png',width=Inches(picture_width-0.5))
paragraph = document.add_paragraph(u'其中二维码、手机闪付、刷卡三种支付方式交易金额随时间变化曲线如下图：')
document.add_picture('支付方式交易金额随时间变化曲线.png',width=Inches(picture_width+0.5))
def to_table(df):
    row = len(df)
    col = len(df.iloc[0])
    tab =document.add_table(rows=row,cols=col,style ='Colorful Grid Accent 4')
    for i in range(row):
        for j in range(col):
            cell=tab.cell(i,j)
            cell.text = str(df.iloc[i,j])          
to_table(machnt_risk)
#document.save(document_dir)
'''
'''
document = Document()
row = len(machnt_risk)
col = len(machnt_risk.iloc[0])
tab =document.add_table(rows=row,cols=col,style ='LightShading-Accent1')
for i in range(row):
    for j in range(col):
        cell=tab.cell(i,j)
        cell.text = str(machnt_risk.iloc[i,j])
document.save(r'C:/工作/典型事件/手机POS交易数据疑似套现/套现分析/分析报告/test.docx')        
'''                       
    
#%%##########################只有二维码交易的商户###############################
def only_QR_func(df):
    QR_list = ['Y','Z']
    if (len(df[df['ext_hce_prod_nm'].isin(QR_list)])==len(df)):
        return True
    else:
        return False   
only_QR = df_sucess.groupby('mchnt_cd').apply(only_QR_func)
df_only_QR = pd.DataFrame({'商户代码':only_QR.index,'是否只有二维码交易':only_QR.values})
df_only_QR = df_only_QR[df_only_QR['是否只有二维码交易']==True]
#%%商户第一笔交易的时间
def first_trade_func(df):
    df = df.sort_index()
    return df['hp_settle_dt_origin'][0]
mchnt_first_trade = df_sucess.groupby('mchnt_cd').apply(first_trade_func)
df_mchnt_first_trade = pd.DataFrame({'商户代码':mchnt_first_trade.index,'首笔交易时间':mchnt_first_trade.values})
df_sended_delay = pd.merge(df_sended,df_mchnt_first_trade,how='left',on='商户代码')
df_sended_delay['暴露时间']=df_sended_delay['报送日期']-df_sended_delay['首笔交易时间']
test_dir = 'C:/工作/典型事件/手机POS交易数据疑似套现/拉卡拉商户交易明细/商户交易日监测报告/delay_time.xlsx'
writer = pd.ExcelWriter(test_dir)
df_sended_delay.to_excel(writer, index=False,encoding='utf-8',sheet_name='Sheet1')
writer.save()
###############################################################################
#%%模板word

from docxtpl import DocxTemplate, RichText, InlineImage
from docx.shared import Mm
import jinja2
import itertools 
context = {} 
# for row, col in itertools.product(machnt_risk.index, machnt_risk.columns):
#     context[f'{row}_{col}'] = df.loc[row, col]
machnt_risk_form = round(machnt_risk,2)
context = machnt_risk_form.to_dict(orient='records')
#context = machnt_risk.set_index('mchnt_cd').T.to_dict('list')
table = {
    'tbl_contents': context
}
'''
table = {
    'user_labels': ['fruit', 'vegetable', 'stone', 'thing'],
    'tbl_contents': [
        {'label': 'yellow', 'cols': ['banana', 'capsicum', 'pyrite', 'taxi']},
        {'label': 'red', 'cols': ['apple', 'tomato', 'cinnabar', 'doubledecker']},
        {'label': 'green', 'cols': ['guava', 'cucumber', 'aventurine', 'card']},
    ],
}
'''

tpl = DocxTemplate(r'C:/工作/典型事件/手机POS交易数据疑似套现/拉卡拉商户交易明细/商户交易日监测报告/商户交易日监测报告_tpl.docx')
rt_date = RichText()
rt_date.add(time_str+'\n', font='方正小标宋简体',size=44)
rt_pargh1 = RichText()
rt_pargh1.add(time_str+'新增商户'+str(mchnt_curr_day)+'家，新增交易笔数'+
             str(trans_num_curr_day)+'笔，新增交易金额'+str(round(trans_at_curr_day,2))+
             '元，贷记卡交易金额占比为'+format(trans_at_loan_curr_day,'.2%')+
             '。本日活跃商户数为：'+str(mchnt_day_live)+
             '，本周活跃商户数为：'+str(mchnt_week_live)+
             '，本月活跃商户数为：'+str(mchnt_month_live)+
             '，本日交易金额为过去30天均值的'+str(round(trans_at_30_ratio,2))+'倍'+
             '，交易笔数为过去30天均值的'+str(round(trans_num_30_ratio,2))+'倍'+
             '，贷记卡交易金额为过去30天均值的'+str(round(loan_trans_at_30_ratio,2))+'倍'+
             '，贷记卡交易笔数为过去30天均值的'+str(round(loan_trans_num_30_ratio,2))+'倍'+   
             '，5000元以上贷记卡交易金额为过去30天均值的'+str(round(loan_trans_5000_at_30_day_ratio,2))+'倍'+
             '，5000元以上贷记卡交易笔数为过去30天均值的'+str(round(loan_trans_5000_num_30_day_ratio,2))+'倍'+
             '，本日uMP上送率为'+str(round(uMP_upload_ratio,4)*100)+'%'+
             '，当日疑似危险商户有'+str(len(machnt_risk))+'家：\n',font='仿宋_GB2312',size=32)

rt_pargh2 = RichText()
rt_pargh2.add('自2020年6月3日以来，手机POS的总交易金额为：'+str(round(trans_at_sum,2))+
              '元,总交易笔数为：'+str(trans_num_sum)+'笔,总商户数量为：'+
              str(mchnt_sum) +'家,总卡片数量为：'+str(card_sum)+'张。\n',
             font='仿宋_GB2312',size=32)
context = {
    'rt_pargh1':rt_pargh1,
    'date': rt_date, 
    'rt_pargh2':rt_pargh2
}

width_pic=130
image = {
    '交易金额随时间变化曲线': InlineImage(tpl, pic_dir+'交易金额随时间变化曲线.png', width=Mm(width_pic)),
    '交易笔数随时间变化曲线': InlineImage(tpl, pic_dir+'交易笔数随时间变化曲线.png', width=Mm(width_pic)),
    '日活跃商户随时间变化趋势': InlineImage(tpl, pic_dir+'日活跃商户随时间变化趋势.png', width=Mm(width_pic)),
    '笔均金额随时间变化曲线': InlineImage(tpl, pic_dir+'笔均金额随时间变化曲线.png', width=Mm(width_pic)),
    '借贷记分布': InlineImage(tpl, pic_dir+'借贷记分布.png', width=Mm(width_pic/1.5)),
    '消费金额区间段分布': InlineImage(tpl, pic_dir+'消费金额区间段分布.png', width=Mm(width_pic)),
    '刷卡与二维码金额比例': InlineImage(tpl, pic_dir+'刷卡与二维码金额比例.png', width=Mm(width_pic/2)),
    '刷卡与二维码笔数比例': InlineImage(tpl, pic_dir+'刷卡与二维码笔数比例.png', width=Mm(width_pic/2)),
    '贷记卡5000元以上交易金额随时间变化曲线': InlineImage(tpl, pic_dir+'贷记卡5000元以上交易金额随时间变化曲线.png', width=Mm(width_pic)),
    '贷记卡5000元以上交易笔数随时间变化曲线': InlineImage(tpl, pic_dir+'贷记卡5000元以上交易笔数随时间变化曲线.png', width=Mm(width_pic)),
    '二维码交易金额区间笔数': InlineImage(tpl, pic_dir+'二维码交易金额区间笔数.png', width=Mm(width_pic)),
    '刷卡交易金额区间笔数': InlineImage(tpl, pic_dir+'刷卡交易金额区间笔数.png', width=Mm(width_pic)),   
    '支付方式交易金额随时间变化曲线': InlineImage(tpl, pic_dir+'支付方式交易金额随时间变化曲线.png', width=Mm(width_pic)),
    '支付方式交易笔均金额随时间变化曲线': InlineImage(tpl, pic_dir+'支付方式交易笔均金额随时间变化曲线.png', width=Mm(width_pic)),
    }

table.update(context)
table.update(image)
jinja_env = jinja2.Environment(autoescape=True)
tpl.render(table, jinja_env)
tpl.save(document_dir)

# %%
